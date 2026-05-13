import re
from copy import deepcopy

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DOC_IN = "/Users/user/Desktop/NIR/_work/nir.docx"
DOC_OUT = "/Users/user/Desktop/NIR/_work/nir_fixed.docx"


def _iter_paragraphs_in_cell(cell):
    for p in cell.paragraphs:
        yield p
    for t in cell.tables:
        yield from iter_paragraphs(t)


def iter_paragraphs(doc_or_table):
    # docx doesn't provide a single iterator across doc+tables; we recurse.
    if hasattr(doc_or_table, "paragraphs") and hasattr(doc_or_table, "tables"):
        for p in doc_or_table.paragraphs:
            yield p
        for t in doc_or_table.tables:
            yield from iter_paragraphs(t)
        return

    # Table
    for row in doc_or_table.rows:
        for cell in row.cells:
            yield from _iter_paragraphs_in_cell(cell)


def run_formatting_sweep(doc: Document):
    # Global defaults (still apply per-run/paragraph to be safe)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)

    # Some Word builds need both name and rFonts to guarantee TNR
    if normal._element.rPr is None:
        normal._element.get_or_add_rPr()
    rfonts = normal._element.rPr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        normal._element.rPr.append(rfonts)
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")

    # Normalize all styles we can access (prevents stray 12pt etc.)
    for st in doc.styles:
        try:
            if hasattr(st, "font") and st.font is not None:
                st.font.name = "Times New Roman"
                st.font.size = Pt(14)
        except Exception:
            pass

    for p in iter_paragraphs(doc):
        # Paragraph formatting
        pf = p.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.line_spacing = 1.5
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.first_line_indent = Cm(1.25)

        # Run formatting cleanup
        for r in p.runs:
            r.font.name = "Times New Roman"
            r.font.size = Pt(14)
            r.font.underline = None
            r.font.strike = False
            r.font.double_strike = False
            # Hard-remove underline/strike from XML (some generators keep it)
            rPr = r._element.rPr
            if rPr is not None:
                for tag in ("w:strike", "w:dstrike", "w:u"):
                    el = rPr.find(qn(tag))
                    if el is not None:
                        rPr.remove(el)


def strip_text_decorations_everywhere(doc: Document):
    """Remove any remaining strike/underline flags from run properties in the XML."""
    body = doc.element.body
    for tag in ("w:strike", "w:dstrike", "w:u"):
        els = body.xpath(f".//{tag}")
        for el in els:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)


def strip_track_changes(doc: Document):
    """
    Remove Word revision markup (w:del / w:ins) so the output does not
    contain struck-through text when exported.
    Strategy:
    - Drop deleted content (<w:del>).
    - Keep inserted content (<w:ins>) by unwrapping it (replace with its children).
    """
    body = doc.element.body

    # Remove all deletions
    dels = body.xpath(".//w:del")
    for d in dels:
        parent = d.getparent()
        if parent is not None:
            parent.remove(d)

    # Unwrap insertions
    ins_list = body.xpath(".//w:ins")
    for ins in ins_list:
        parent = ins.getparent()
        if parent is None:
            continue
        idx = parent.index(ins)
        # insert children in place of <w:ins>
        for child in list(ins):
            parent.insert(idx, deepcopy(child))
            idx += 1
        parent.remove(ins)


def fix_vvedenie_wording(doc: Document):
    # Replace the required lead-in for "Актуальность"
    for p in iter_paragraphs(doc):
        text = p.text
        if "Актуальность темы определяется тем, что" in text:
            # Keep the rest of sentence
            new_text = text.replace(
                "Актуальность темы определяется тем, что",
                "Актуальность работы состоит в том, что",
            )
            # Rewrite paragraph preserving runs isn't worth it here; rebuild runs.
            for r in p.runs:
                r.text = ""
            p.add_run(new_text)
            break


def fix_title_quotes(doc: Document):
    # On title page the theme line contains "на тему:"; wrap the theme in quotes
    # and remove "на тему:" if present.
    theme_re = re.compile(r"^\s*на тему:\s*(.+)\s*$", re.IGNORECASE)
    for p in iter_paragraphs(doc):
        m = theme_re.match(p.text.strip())
        if not m:
            continue
        theme = m.group(1).strip().strip("«»\"")
        fixed = f"«{theme}»"
        for r in p.runs:
            r.text = ""
        p.text = fixed
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Title requirement: 14pt center; rest handled by formatting sweep.
        break


def fix_figure_captions(doc: Document):
    # Convert standalone caption lines like "рисунок 2.1 карта ..." to "Рисунок 2.1 – ..."
    cap_re = re.compile(r"^\s*рисунок\s+(\d+(?:\.\d+)*)\s+(.+?)\s*$", re.IGNORECASE)
    for p in iter_paragraphs(doc):
        m = cap_re.match(p.text)
        if not m:
            continue
        num = m.group(1)
        title = m.group(2).strip()
        fixed = f"Рисунок {num} – {title}"
        for r in p.runs:
            r.text = ""
        p.text = fixed
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _get_or_add_sectPr(el_body):
    sectPr = el_body.xpath("./w:sectPr")
    if sectPr:
        return sectPr[0]
    sectPr = OxmlElement("w:sectPr")
    el_body.append(sectPr)
    return sectPr


def remove_page_borders_from_all_sections_except_first(doc: Document):
    """
    Keep page border only on the first section.
    If the document currently has only one section, we can't reliably apply a border
    to "only title page" without restructuring; in that case we remove borders entirely
    (this still satisfies 'no рамка далее по тексту'), and the title border can be added manually.
    """
    if len(doc.sections) <= 1:
        # Remove any page borders from the only section
        sectPr = doc.sections[0]._sectPr
        pgBorders = sectPr.find(qn("w:pgBorders"))
        if pgBorders is not None:
            sectPr.remove(pgBorders)
        return

    # Remove from all sections after first
    for s in doc.sections[1:]:
        sectPr = s._sectPr
        pgBorders = sectPr.find(qn("w:pgBorders"))
        if pgBorders is not None:
            sectPr.remove(pgBorders)


def main():
    doc = Document(DOC_IN)

    strip_track_changes(doc)
    fix_title_quotes(doc)
    fix_vvedenie_wording(doc)
    fix_figure_captions(doc)
    run_formatting_sweep(doc)
    strip_text_decorations_everywhere(doc)
    remove_page_borders_from_all_sections_except_first(doc)

    doc.save(DOC_OUT)


if __name__ == "__main__":
    main()

