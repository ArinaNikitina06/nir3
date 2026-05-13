#!/usr/bin/env python3
"""
Приведение НИР к требованиям чек-листа МИСИС (шрифт, интервалы, титул, рамка только на 1-й стр.).
Исходник не перезаписывается — создаётся новый .docx рядом.
"""
from __future__ import annotations

import re
import shutil
import sys
import warnings
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

SRC = Path(__file__).resolve().parent / "МИСТ-24-3-3_Никитина_без_1_приложения_3сем.docx"
OUT = Path(__file__).resolve().parent / "МИСТ-24-3-3_Никитина_НИР_оформлено.docx"

TITLE_LAST = 28  # включая «Москва 2026»
TOC_HEADER = 29
TOC_BLOCK_END = 44  # до «ВВЕДЕНИЕ»


def set_run_font(run, size_pt: int = 14, bold=None, italic=None) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    for k in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{k}"), "Times New Roman")


def paragraph_clear_set(p, text: str, *, size_pt=14, bold=None, italic=None) -> None:
    p.clear()
    r = p.add_run(text)
    set_run_font(r, size_pt=size_pt, bold=bold, italic=italic)


def add_page_border_first_page_only(section) -> None:
    """Двойная рамка страницы только на первой странице раздела (как в Word «только на первой»)."""
    sectPr = section._sectPr
    for el in list(sectPr):
        if el.tag == qn("w:pgBorders"):
            sectPr.remove(el)
    pg = OxmlElement("w:pgBorders")
    pg.set(qn("w:offsetFrom"), "page")
    pg.set(qn("w:display"), "firstPage")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "double")
        b.set(qn("w:sz"), "18")  # толщина линии (в 1/8 pt)
        b.set(qn("w:space"), "28")  # отступ рамки от края листа, twips-подобная величина в Word
        b.set(qn("w:color"), "000000")
        pg.append(b)
    sectPr.append(pg)


def add_paragraph_bottom_double_line(p) -> None:
    """Двойная горизонтальная линия под абзацем (разделитель на титуле)."""
    pPr = p._element.get_or_add_pPr()
    for el in list(pPr):
        if el.tag == qn("w:pBdr"):
            pPr.remove(el)
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "double")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)


def apply_paragraph_base(
    p,
    *,
    align,
    first_indent_cm=None,
    line_multiple=1.5,
    space_before_pt=0,
    space_after_pt=0,
) -> None:
    pf = p.paragraph_format
    pf.alignment = align
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_multiple
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    if first_indent_cm is None:
        pf.first_line_indent = None
    else:
        pf.first_line_indent = Cm(first_indent_cm)


def style_update_heading(doc, name: str) -> None:
    st = doc.styles[name]
    st.font.name = "Times New Roman"
    st.font.size = Pt(14)
    st.font.bold = True
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = None


def style_update_normal(doc) -> None:
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(14)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def format_title_page(doc) -> None:
    # Министерство / ФГАОУ
    for idx in (0, 2):
        p = doc.paragraphs[idx]
        paragraph_clear_set(p, p.text.strip(), size_pt=14, bold=False, italic=False)
        apply_paragraph_base(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_cm=None)
        for r in p.runs:
            set_run_font(r, 14, bold=False, italic=False)

    # Университет — жирный, как в шаблоне
    p3 = doc.paragraphs[3]
    paragraph_clear_set(
        p3,
        "«НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ ТЕХНОЛОГИЧЕСКИЙ УНИВЕРСИТЕТ «МИСИС»",
        size_pt=14,
        bold=True,
        italic=False,
    )
    apply_paragraph_base(p3, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_cm=None)
    add_paragraph_bottom_double_line(p3)

    # Институт / кафедра / направление — курсив, прописные
    for idx in (6, 7, 8):
        p = doc.paragraphs[idx]
        t = p.text.strip()
        paragraph_clear_set(p, t, size_pt=14, bold=False, italic=True)
        apply_paragraph_base(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_cm=None)
        for r in p.runs:
            set_run_font(r, 14, bold=False, italic=True)

    # «ОТЧЁТ» — чуть крупнее для читаемости титула
    p12 = doc.paragraphs[12]
    paragraph_clear_set(p12, "ОТЧЁТ", size_pt=16, bold=True, italic=False)
    apply_paragraph_base(p12, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_cm=None)

    p13 = doc.paragraphs[13]
    paragraph_clear_set(
        p13, "по научно-исследовательской работе", size_pt=14, bold=True, italic=False
    )
    apply_paragraph_base(p13, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_cm=None)

    p14 = doc.paragraphs[14]
    topic = (
        "на тему: Исследование UX-паттернов и принципов проектирования интерфейсов "
        "персонализированных рекомендательных систем для веб-приложений онлайн-обучения"
    )
    paragraph_clear_set(p14, topic, size_pt=14, bold=True, italic=False)
    apply_paragraph_base(p14, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_cm=None)

    # Правый блок «Выполнил»
    p17 = doc.paragraphs[17]
    paragraph_clear_set(p17, "Выполнил:", size_pt=14, bold=True, italic=False)
    apply_paragraph_base(
        p17,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        first_indent_cm=None,
    )
    p17.paragraph_format.left_indent = Cm(8.5)

    pairs = [
        (18, "Студент группы МИСТ-24-3-3", False, True),
        (19, "Никитина Арина Олеговна", False, True),
        (20, "Научный руководитель НИР", False, True),
        (21, "Лозинская Мария Алексеевна", False, True),
    ]
    for idx, text, bold, italic in pairs:
        p = doc.paragraphs[idx]
        paragraph_clear_set(p, text, size_pt=14, bold=bold, italic=italic)
        apply_paragraph_base(
            p,
            align=WD_ALIGN_PARAGRAPH.LEFT,
            first_indent_cm=None,
        )
        p.paragraph_format.left_indent = Cm(8.5)

    p28 = doc.paragraphs[28]
    paragraph_clear_set(p28, "Москва 2026", size_pt=14, bold=False, italic=False)
    apply_paragraph_base(p28, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_cm=None)

    # Пустые строки титула — без лишних отступов
    for idx in range(0, TITLE_LAST + 1):
        if idx in (
            0,
            2,
            3,
            6,
            7,
            8,
            12,
            13,
            14,
            17,
            18,
            19,
            20,
            21,
            28,
        ):
            continue
        p = doc.paragraphs[idx]
        if not p.text.strip():
            apply_paragraph_base(
                p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_cm=None
            )
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)


def format_toc_block(doc) -> None:
    ph = doc.paragraphs[TOC_HEADER]
    paragraph_clear_set(ph, "СОДЕРЖАНИЕ", size_pt=14, bold=True, italic=False)
    apply_paragraph_base(ph, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent_cm=None)
    for i in range(TOC_HEADER + 1, TOC_BLOCK_END + 1):
        p = doc.paragraphs[i]
        apply_paragraph_base(
            p,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            first_indent_cm=None,
        )
        for r in p.runs:
            set_run_font(r, 14)


def is_figure_caption(text: str) -> bool:
    t = text.strip()
    return bool(re.match(r"^рисунок\s+[\d\.]+", t, flags=re.I))


def format_body(doc) -> None:
    for i, p in enumerate(doc.paragraphs):
        if i <= TOC_BLOCK_END:
            continue
        name = p.style.name
        raw = p.text
        if not raw.strip():
            apply_paragraph_base(
                p,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_indent_cm=None,
            )
            continue

        if name == "Heading 1":
            apply_paragraph_base(
                p,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                first_indent_cm=None,
            )
            for r in p.runs:
                set_run_font(r, 14, bold=True, italic=False)
            continue

        if name in ("Heading 2", "Heading 3"):
            apply_paragraph_base(
                p,
                align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                first_indent_cm=None,
            )
            for r in p.runs:
                set_run_font(r, 14, bold=True, italic=False)
            continue

        if is_figure_caption(raw):
            # Чек-лист: «Рисунок N – Название», подпись снизу, 14 TNR
            t = raw.strip()
            t = re.sub(
                r"^рисунок\s+([\d\.]+)\s*(.*)$",
                r"Рисунок \1 – \2",
                t,
                flags=re.I,
            ).strip()
            paragraph_clear_set(p, t, size_pt=14, bold=False, italic=False)
            apply_paragraph_base(
                p,
                align=WD_ALIGN_PARAGRAPH.CENTER,
                first_indent_cm=None,
            )
            continue

        apply_paragraph_base(
            p,
            align=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_indent_cm=1.25,
        )
        for r in p.runs:
            if r.text.startswith(" ") or r.text.endswith(" "):
                r.text = r.text.strip(" ")
            set_run_font(r, 14, bold=r.bold, italic=r.italic)


def main() -> int:
    if not SRC.exists():
        print("Нет файла:", SRC, file=sys.stderr)
        return 1
    warnings.filterwarnings(
        "ignore",
        message="style lookup by style_id is deprecated",
        category=UserWarning,
    )
    shutil.copyfile(SRC, OUT)
    doc = Document(str(OUT))

    style_update_normal(doc)
    for hn in ("Heading 1", "Heading 2", "Heading 3"):
        if hn in [s.name for s in doc.styles]:
            style_update_heading(doc, hn)

    add_page_border_first_page_only(doc.sections[0])
    format_title_page(doc)
    format_toc_block(doc)
    format_body(doc)

    doc.save(str(OUT))
    print("Записано:", OUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
